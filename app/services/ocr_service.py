import base64
import io
from docx import Document
from docx.shared import Pt
import fitz

from app.services.ollama import ollama_service

PAGES_PER_CHUNK = 1
DPI = 150

EXTRACTION_PROMPT = """Actúa como un ingeniero de datos experto en arquitecturas RAG (Retrieval-Augmented Generation) y extracción de información. Tu tarea es analizar la imagen adjunta y extraer todo su contenido en formato texto, estructurándolo de la forma más óptima para ser indexado en una base de datos vectorial.

Sigue estrictamente estas directrices de formato y extracción:

1. Estructura Jerárquica y Semántica:
   - Utiliza Markdown puro.
   - Usa encabezados (`#`, `##`, `###`) para reflejar fielmente la jerarquía original del documento. El sistema RAG utilizará estos encabezados para mantener el contexto durante el 'chunking'.

2. Extracción Exacta de Texto:
   - Transcribe el texto de forma fiel y completa.
   - No resumas, no parafrasees y no omitas información.
   - Corrige silenciosamente los errores obvios de escaneo (caracteres basura generados por el OCR), pero mantén la terminología técnica o legal exacta.

3. Procesamiento de Tablas (Regla Crítica):
   - Tablas Simples: Si la tabla es plana y sencilla, conviértela en una tabla estándar de Markdown (utilizando `|` y `-`).
   - Tablas Complejas (Anexos, instructivos con celdas combinadas): Si la tabla es compleja, usar Markdown tradicional romperá el contexto al fragmentarse (chunking). En su lugar, transforma la tabla en una estructura de lista semántica de clave-valor por cada fila o ítem. Utiliza este formato:
     **[Título de la Tabla o Sección]**
     * Ítem/Fila 1:
       - [Nombre de Columna 1]: [Valor extraído]
       - [Nombre de Columna 2]: [Valor extraído]
       - [Nombre de Columna 3]: [Valor extraído]

4. Listas y Enumeraciones:
   - Mantén estrictamente la numeración original (letras o números) o utiliza viñetas estándar de Markdown (`*` o `-`) para facilitar la separación de párrafos.

5. Elementos No Textuales Relevantes:
   - Si la imagen contiene sellos, firmas, marcas de agua o diagramas que aportan validez o contexto (ej. "Sello de la Aduana Nacional", "Firma de la Máxima Autoridad"), descríbelos brevemente entre corchetes. Ejemplo: `[Elemento visual: Firma y sello del Gerente Nacional de Normas]`.

No incluyas saludos, introducciones ni conclusiones en tu respuesta. Devuelve única y exclusivamente el texto procesado en formato Markdown, listo para ser ingerido por el pipeline de datos."""


async def extract_pdf_to_word(pdf_bytes: bytes, filename: str, ocr_model: str, on_progress=None) -> bytes:
    print(f"\n{'='*60}")
    print(f"  🔍 OCR EXTRACTOR — Iniciando")
    print(f"  📄 Archivo: {filename}")
    print(f"  🤖 Modelo: {ocr_model}")
    print(f"  📐 DPI: {DPI} | Paginas por chunk: {PAGES_PER_CHUNK}")
    print(f"{'='*60}")

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    total_pages = len(doc)

    if total_pages == 0:
        doc.close()
        raise ValueError("El PDF no tiene paginas")

    print(f"  📊 Total de paginas: {total_pages}")

    chunks = [(i, min(i + PAGES_PER_CHUNK, total_pages)) for i in range(0, total_pages, PAGES_PER_CHUNK)]
    total_chunks = len(chunks)
    all_text: list[str] = []

    print(f"  📦 Total de chunks: {total_chunks}")
    print(f"{'='*60}\n")

    for chunk_idx, (start, end) in enumerate(chunks):
        images: list[str] = []
        for page_num in range(start, end):
            page = doc[page_num]
            pix = page.get_pixmap(dpi=DPI)
            img_bytes = pix.tobytes("png")
            img_b64 = base64.b64encode(img_bytes).decode("utf-8")
            images.append(img_b64)
            print(f"  📸 Pagina {page_num + 1}/{total_pages} renderizada ({pix.width}x{pix.height}, {len(img_bytes)//1024}KB)")

        print(f"  🤖 Enviando chunk {chunk_idx + 1}/{total_chunks} (paginas {start+1}-{end}) a {ocr_model}...")
        text = await ollama_service.chat_with_images(images, EXTRACTION_PROMPT, model=ocr_model)
        text = text.strip()
        all_text.append(text)

        text_preview = text[:100].replace("\n", " ")
        print(f"  ✅ Chunk {chunk_idx + 1}/{total_chunks} completado — {len(text)} caracteres extraidos")
        print(f"     Preview: {text_preview}...")

        if on_progress:
            await on_progress(chunk_idx + 1, total_chunks, f"Chunk {chunk_idx + 1}/{total_chunks} procesado")

    doc.close()
    print(f"\n  📝 Generando documento Word...")

    docx_bytes = _generate_docx(all_text, filename)

    total_chars = sum(len(t) for t in all_text)
    print(f"  ✅ Word generado: {total_chars} caracteres totales, {len(all_text)} paginas")
    print(f"  📦 Tamano del .docx: {len(docx_bytes)//1024}KB")
    print(f"{'='*60}\n")

    return docx_bytes


def _generate_docx(text_chunks: list[str], original_filename: str) -> bytes:
    docx_doc = Document()

    style = docx_doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for i, text in enumerate(text_chunks):
        if i > 0:
            docx_doc.add_page_break()

        paragraphs = text.split("\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            if _is_heading(para_text):
                docx_doc.add_heading(para_text, level=2)
            else:
                docx_doc.add_paragraph(para_text)

    buffer = io.BytesIO()
    docx_doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _is_heading(text: str) -> bool:
    if len(text) > 100:
        return False
    upper_ratio = sum(1 for c in text if c.isupper()) / max(len(text), 1)
    return upper_ratio > 0.6 or text.endswith(":") and len(text) < 60
